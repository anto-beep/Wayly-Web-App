"""
Document extraction — single entry point that routes any supported upload
(PDF text/scanned, DOCX, TXT, JPG/PNG/HEIC/WEBP) into plain text, falling
back to Claude vision for image and scanned-PDF inputs.

Returns: (text, input_method, document_pages, parsing_warnings)

Supported extensions:
  Text path: .pdf (selectable), .docx, .txt
  Vision path: .pdf (scanned), .jpg, .jpeg, .png, .heic, .heif, .webp
"""
from __future__ import annotations
import asyncio
import base64
import io
import logging
import os
import re
from typing import Tuple, List, Optional, Dict, Any

logger = logging.getLogger(__name__)

# Lazy imports to avoid forcing heavy deps at module-load time.
_HEIF_REGISTERED = False

# Image quality thresholds (tuned for phone photos of A4 statements)
_QUALITY_THRESHOLDS = {
    "brightness_low": 60,        # mean luminance — below = too dark
    "brightness_high": 245,      # above = washed out / overexposed
    "blur_min_good": 150.0,      # Laplacian variance — above = sharp
    "blur_min_fair": 60.0,       # below = poor blur
    "skew_warn_deg": 4.0,        # rotation > this triggers an info note
    "skew_correct_deg": 1.0,     # skews above this trigger auto-rotation
    "resolution_min_short": 900, # short side in pixels — below = low-res
    "resolution_low_short": 600, # short side below this = poor (likely unreadable)
    "blank_page_var": 8.0,       # pixel-stddev below this = effectively blank
    "blank_page_min_brightness": 200,  # blank pages are also bright (white paper)
}


ALLOWED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".txt", ".csv",
    ".jpg", ".jpeg", ".png", ".heic", ".heif", ".webp",
}

EXT_TO_MIME = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".txt": "text/plain",
    ".csv": "text/csv",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".heic": "image/heic",
    ".heif": "image/heif",
    ".webp": "image/webp",
}

# Format-specific size limits (bytes)
MAX_BYTES = {
    ".pdf": 20 * 1024 * 1024,
    ".docx": 10 * 1024 * 1024,
    ".doc": 10 * 1024 * 1024,
    ".txt": 5 * 1024 * 1024,
    ".csv": 10 * 1024 * 1024,
    ".jpg": 10 * 1024 * 1024,
    ".jpeg": 10 * 1024 * 1024,
    ".png": 10 * 1024 * 1024,
    ".heic": 10 * 1024 * 1024,
    ".heif": 10 * 1024 * 1024,
    ".webp": 10 * 1024 * 1024,
}


# Magic-byte signatures for the formats we accept.
def _verify_magic_bytes(ext: str, raw: bytes) -> bool:
    if not raw:
        return False
    head = raw[:16]
    if ext == ".pdf":
        return head[:4] == b"%PDF"
    if ext in (".jpg", ".jpeg"):
        return head[:2] == b"\xff\xd8"
    if ext == ".png":
        return head[:8] == b"\x89PNG\r\n\x1a\n"
    if ext == ".webp":
        return head[:4] == b"RIFF" and head[8:12] == b"WEBP"
    if ext in (".heic", ".heif"):
        # ISO base media file format with ftypheic / ftypmif1 / ftypheix at byte 4
        return head[4:8] == b"ftyp"
    if ext == ".docx":
        # DOCX is a ZIP — first 2 bytes "PK"
        return head[:2] == b"PK"
    if ext == ".doc":
        return head[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
    # Text formats — accept anything decodable
    if ext in (".txt", ".csv"):
        return True
    return True


def get_extension(filename: str) -> str:
    name = (filename or "").lower()
    m = re.search(r"\.[a-z0-9]{2,5}$", name)
    return m.group(0) if m else ""


class UnsupportedFormatError(Exception):
    pass


class FileTooLargeError(Exception):
    def __init__(self, ext: str, limit_bytes: int):
        self.ext = ext
        self.limit_bytes = limit_bytes
        super().__init__(f"{ext} exceeds limit of {limit_bytes} bytes")


class CorruptFileError(Exception):
    pass


class PasswordProtectedError(Exception):
    pass


def validate_upload(filename: str, raw: bytes) -> str:
    """Returns the validated extension. Raises typed exceptions on failure."""
    ext = get_extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise UnsupportedFormatError(ext or "(no extension)")
    limit = MAX_BYTES.get(ext, 10 * 1024 * 1024)
    if len(raw) > limit:
        raise FileTooLargeError(ext, limit)
    if not _verify_magic_bytes(ext, raw):
        raise CorruptFileError(f"Magic-byte check failed for {ext}")
    return ext


# ───────── Image quality assessment (OpenCV) ─────────
#
# Detects brightness, blur, skew, low resolution, and blank pages on phone
# photos of paper statements. Returns a structured dict + a list of
# user-facing warnings. Auto-rotates skewed images so Claude vision sees
# upright text.

def _cv2():
    """Lazy import of opencv. Returns None if unavailable so quality
    assessment degrades gracefully without breaking the pipeline."""
    try:
        import cv2  # type: ignore
        return cv2
    except ImportError:
        logger.warning("opencv-python-headless not installed — image quality assessment skipped")
        return None


def _np():
    try:
        import numpy as np  # type: ignore
        return np
    except ImportError:
        return None


def _estimate_skew_angle(gray) -> float:
    """Returns estimated skew in degrees (positive = clockwise lean).
    Uses HoughLinesP on edge map for robustness against axis-aligned content."""
    cv2 = _cv2()
    np = _np()
    if cv2 is None or np is None:
        return 0.0
    try:
        h, w = gray.shape
        # Downscale for speed and noise reduction
        scale = 1000.0 / max(h, w) if max(h, w) > 1000 else 1.0
        if scale < 1.0:
            small = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_AREA)
        else:
            small = gray
        edges = cv2.Canny(small, 50, 150, apertureSize=3)
        min_line_len = max(50, int(0.25 * min(small.shape)))
        lines = cv2.HoughLinesP(
            edges, rho=1, theta=np.pi / 360,
            threshold=100,
            minLineLength=min_line_len,
            maxLineGap=20,
        )
        if lines is None or len(lines) == 0:
            return 0.0
        angles = []
        for x1, y1, x2, y2 in lines[:, 0, :]:
            dx = x2 - x1
            dy = y2 - y1
            if dx == 0:
                continue
            angle = np.degrees(np.arctan2(dy, dx))
            # Normalise to nearest horizontal: collapse vertical lines
            if angle > 45:
                angle -= 90
            elif angle < -45:
                angle += 90
            # Only consider near-horizontal lines (table rows / text baselines)
            if abs(angle) <= 30:
                angles.append(angle)
        if len(angles) < 5:
            return 0.0
        # Use median (robust to outliers)
        return float(np.median(angles))
    except Exception as e:
        logger.warning("Skew estimate failed: %s", e)
        return 0.0


def assess_image_quality(pil_img) -> Dict[str, Any]:
    """Returns {brightness, blur_score, skew_angle, width, height, rating,
    is_blank, warnings[]}. Does NOT mutate the image."""
    cv2 = _cv2()
    np = _np()
    if cv2 is None or np is None:
        return {
            "brightness": None, "blur_score": None, "skew_angle": 0.0,
            "width": pil_img.size[0], "height": pil_img.size[1],
            "rating": "unknown", "is_blank": False, "warnings": [],
        }
    try:
        arr = np.array(pil_img.convert("RGB"))
        gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
    except Exception as e:
        logger.warning("Quality assessment failed: %s", e)
        return {
            "brightness": None, "blur_score": None, "skew_angle": 0.0,
            "width": pil_img.size[0], "height": pil_img.size[1],
            "rating": "unknown", "is_blank": False, "warnings": [],
        }
    h, w = gray.shape
    brightness = float(gray.mean())
    pixel_std = float(gray.std())
    blur_score = float(cv2.Laplacian(gray, cv2.CV_64F).var())
    skew_angle = _estimate_skew_angle(gray)
    short_side = min(w, h)

    th = _QUALITY_THRESHOLDS
    warnings: list[str] = []
    rating = "good"
    is_blank = (pixel_std < th["blank_page_var"]) and (brightness > th["blank_page_min_brightness"])

    if is_blank:
        return {
            "brightness": round(brightness, 1),
            "blur_score": round(blur_score, 1),
            "skew_angle": round(skew_angle, 2),
            "width": w, "height": h,
            "rating": "blank", "is_blank": True, "warnings": ["Page appears blank — skipped."],
        }

    if brightness < th["brightness_low"]:
        warnings.append("The photo is quite dark — text may not extract cleanly. Try better lighting.")
        rating = "poor"
    elif brightness > th["brightness_high"]:
        warnings.append("The photo looks overexposed (washed out) — some text may be lost.")
        if rating == "good":
            rating = "fair"

    if blur_score < th["blur_min_fair"]:
        warnings.append("The image is blurry. Hold the camera steady and tap to focus before the shot.")
        rating = "poor"
    elif blur_score < th["blur_min_good"]:
        if rating == "good":
            rating = "fair"

    if abs(skew_angle) > th["skew_warn_deg"]:
        warnings.append(
            f"The page is tilted about {abs(skew_angle):.0f}°. We've auto-rotated it but accuracy may suffer — "
            "next time, photograph the page flat-on."
        )

    if short_side < th["resolution_low_short"]:
        warnings.append("Image resolution is very low — text may be unreadable. Try a higher-quality scan.")
        rating = "poor"
    elif short_side < th["resolution_min_short"]:
        if rating == "good":
            rating = "fair"
        warnings.append("Image resolution is on the low side — consider a higher-quality scan if results look off.")

    return {
        "brightness": round(brightness, 1),
        "blur_score": round(blur_score, 1),
        "skew_angle": round(skew_angle, 2),
        "width": w, "height": h,
        "rating": rating,
        "is_blank": False,
        "warnings": warnings,
    }


def auto_rotate(pil_img, skew_angle: float):
    """Rotates the image to compensate for detected skew. Returns the rotated
    PIL image (or the original if skew is below the correction threshold)."""
    if abs(skew_angle) < _QUALITY_THRESHOLDS["skew_correct_deg"]:
        return pil_img
    cv2 = _cv2()
    np = _np()
    if cv2 is None or np is None:
        return pil_img
    try:
        from PIL import Image
        arr = np.array(pil_img.convert("RGB"))
        h, w = arr.shape[:2]
        M = cv2.getRotationMatrix2D((w / 2, h / 2), skew_angle, 1.0)
        rotated = cv2.warpAffine(
            arr, M, (w, h),
            flags=cv2.INTER_CUBIC,
            borderMode=cv2.BORDER_REPLICATE,
        )
        return Image.fromarray(rotated)
    except Exception as e:
        logger.warning("Auto-rotate failed: %s", e)
        return pil_img


def _prepare_image_for_vision(pil_img) -> Tuple[Any, Dict[str, Any]]:
    """Assess + auto-rotate. Returns (prepared_pil_img, quality_dict)."""
    quality = assess_image_quality(pil_img)
    if not quality.get("is_blank") and abs(quality.get("skew_angle") or 0.0) >= _QUALITY_THRESHOLDS["skew_correct_deg"]:
        pil_img = auto_rotate(pil_img, quality["skew_angle"])
    return pil_img, quality


# ───────── Text extraction paths ─────────

def _extract_pdf_selectable_text(raw: bytes) -> Tuple[str, int]:
    """Returns (text, page_count). Raises PasswordProtectedError on locked PDFs."""
    try:
        import pdfplumber
    except ImportError:
        # Fall back to pypdf only.
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        if getattr(reader, "is_encrypted", False):
            raise PasswordProtectedError("PDF is password-protected")
        text = "\n".join((p.extract_text() or "") for p in reader.pages)
        return text, len(reader.pages)
    try:
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            pages = pdf.pages
            page_count = len(pages)
            chunks = []
            for p in pages:
                try:
                    t = p.extract_text() or ""
                except Exception:
                    t = ""
                if t:
                    chunks.append(t)
            return "\n".join(chunks), page_count
    except Exception as e:
        msg = str(e).lower()
        if "password" in msg or "encrypted" in msg:
            raise PasswordProtectedError("PDF is password-protected") from e
        # Fall back to pypdf
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        if getattr(reader, "is_encrypted", False):
            raise PasswordProtectedError("PDF is password-protected")
        text = "\n".join((p.extract_text() or "") for p in reader.pages)
        return text, len(reader.pages)


def _extract_docx_text(raw: bytes) -> str:
    """Returns text. Tables are flattened tab-separated to preserve column structure."""
    try:
        import docx as _docx
    except ImportError as e:
        raise CorruptFileError(f"python-docx missing: {e}")
    try:
        doc = _docx.Document(io.BytesIO(raw))
    except Exception as e:
        raise CorruptFileError(f"DOCX parse failed: {e}") from e
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _extract_txt(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise CorruptFileError("Could not decode text file with common encodings")


# ───────── Vision path ─────────

def _ensure_heif_registered():
    global _HEIF_REGISTERED
    if _HEIF_REGISTERED:
        return
    try:
        from pillow_heif import register_heif_opener
        register_heif_opener()
        _HEIF_REGISTERED = True
    except Exception as e:
        logger.warning("HEIF support unavailable: %s", e)


def _image_to_base64_jpeg(raw: bytes, ext: str) -> Tuple[str, Dict[str, Any]]:
    """Decodes the image, runs quality assessment, auto-rotates if skewed,
    and returns (base64_jpeg, quality_dict)."""
    _ensure_heif_registered()
    from PIL import Image
    try:
        img = Image.open(io.BytesIO(raw))
        img.load()  # force-decode so EXIF orientation can apply
    except Exception as e:
        raise CorruptFileError(f"Could not open {ext} image: {e}") from e
    # Apply EXIF orientation so portrait phone photos aren't sideways
    try:
        from PIL import ImageOps
        img = ImageOps.exif_transpose(img)
    except Exception:
        pass
    # Normalise mode for JPEG output
    if img.mode in ("RGBA", "LA", "P"):
        img = img.convert("RGB")
    img, quality = _prepare_image_for_vision(img)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=92)
    return base64.b64encode(buf.getvalue()).decode("ascii"), quality


def _pdf_to_image_pages_b64(raw: bytes, max_pages: int = 8) -> List[Dict[str, Any]]:
    """Converts a scanned PDF into a list of {b64, quality, page_num} records,
    one per non-blank page. Limits to max_pages to keep cost predictable.
    Each page is run through the same quality + auto-rotate pipeline as
    standalone images."""
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        # pdf2image isn't available in this env — bail gracefully.
        raise CorruptFileError(
            "Scanned PDF processing isn't available in this environment. "
            "Try uploading the statement as a JPG photo or pasting the text."
        )
    try:
        images = convert_from_bytes(raw, dpi=200, fmt="jpeg", thread_count=2)
    except Exception as e:
        raise CorruptFileError(f"PDF to image conversion failed: {e}") from e
    out: list[Dict[str, Any]] = []
    for i, img in enumerate(images[:max_pages]):
        if img.mode in ("RGBA", "LA", "P"):
            img = img.convert("RGB")
        prepared, quality = _prepare_image_for_vision(img)
        if quality.get("is_blank"):
            # Drop blank pages but record their existence in the result
            out.append({"b64": None, "quality": quality, "page_num": i + 1, "skipped": True})
            continue
        buf = io.BytesIO()
        prepared.save(buf, format="JPEG", quality=88)
        out.append({
            "b64": base64.b64encode(buf.getvalue()).decode("ascii"),
            "quality": quality,
            "page_num": i + 1,
            "skipped": False,
        })
    return out


VISION_EXTRACTION_PROMPT = """You are reading a photographed or scanned Australian Support at Home monthly statement. Extract every piece of information visible in this image as plain text — preserving the original column structure of any tables.

Pay particular attention to:

TABLES: Statement line items are usually in a table with columns for date, service description, service code, hours, rate per hour, gross amount, participant contribution, and government paid amount. The columns may not be clearly separated — use context to determine which number belongs to which column. Output each table row on its own line with values tab-separated.

DOLLAR AMOUNTS: Read every dollar figure precisely. $1,047.00 and $1,047.50 are different. Never round or approximate dollar amounts. If a figure is genuinely unclear, transcribe it as "[unclear: best-guess]".

DATES: Read dates exactly as written. Do not infer missing dates.

PROVIDER NAME AND ABN: Usually in the header section. Extract exactly as printed including any formatting.

STREAM LABELS: Look for the words Clinical, Independence, and Everyday Living as section headers.

HANDWRITTEN ANNOTATIONS: If any text appears handwritten rather than printed, prefix the line with "[HANDWRITTEN]:".

If a portion of the document is genuinely unreadable, transcribe what you can and add a note like "[unclear region]" inline.

Return ONLY the extracted text content — do NOT add commentary, headings of your own, or markdown. Just the statement as you would re-type it from the image, line by line.
"""


async def _vision_extract_image(image_b64: str, label: str = "image") -> str:
    """Sends an image to Claude vision and returns extracted text."""
    from emergentintegrations.llm.chat import LlmChat, UserMessage, ImageContent
    api_key = os.environ.get("EMERGENT_LLM_KEY")
    if not api_key:
        raise CorruptFileError("Vision processing unavailable — EMERGENT_LLM_KEY not configured")
    chat = LlmChat(
        api_key=api_key,
        session_id=f"vision-{label}",
        system_message=VISION_EXTRACTION_PROMPT,
    ).with_model("anthropic", "claude-sonnet-4-5-20250929").with_params(max_tokens=4000)
    image_content = ImageContent(image_base64=image_b64)
    msg = UserMessage(text="Read this Support at Home statement.", file_contents=[image_content])
    try:
        result = await chat.send_message(msg)
    except Exception as e:
        raise CorruptFileError(f"Vision extraction failed: {e}") from e
    return str(result or "")


# ───────── Public API ─────────

async def extract_document(filename: str, raw: bytes) -> Tuple[str, str, int, List[str]]:
    """Routes a validated upload to the right extraction path.

    Returns (text, input_method, page_count, parsing_warnings).

    input_method values:
      text_paste, text_file, word_document, pdf_text, pdf_scanned,
      image_vision, email_attachment
    """
    ext = validate_upload(filename, raw)
    warnings: list[str] = []

    if ext == ".txt":
        return _extract_txt(raw), "text_file", 1, warnings

    if ext == ".csv":
        return _extract_txt(raw), "text_file", 1, warnings

    if ext == ".docx":
        return _extract_docx_text(raw), "word_document", 1, warnings

    if ext == ".doc":
        # Legacy DOC — Phase 2 will add LibreOffice. For now route to the vision
        # path by complaining loudly so the user can re-export as DOCX or PDF.
        raise UnsupportedFormatError(
            "Legacy .doc files aren't supported yet — open the file in Word and "
            "save as .docx or PDF, then upload again."
        )

    if ext == ".pdf":
        text, page_count = _extract_pdf_selectable_text(raw)
        # Decide if PDF has enough selectable text or needs vision.
        cleaned = (text or "").strip()
        keywords = ("Support at Home", "participant", "Classification", "stream", "$", "quarterly")
        has_signal = any(k in cleaned for k in keywords)
        if len(cleaned) > 500 and has_signal:
            return cleaned, "pdf_text", page_count, warnings
        # Fall through to scanned-PDF vision path.
        try:
            page_records = _pdf_to_image_pages_b64(raw, max_pages=8)
        except CorruptFileError:
            # No vision possible — return whatever text we got. Better than 0.
            if cleaned:
                warnings.append("PDF text was sparse and scanned-PDF processing is unavailable in this environment. Result may be incomplete.")
                return cleaned, "pdf_text", page_count, warnings
            raise
        usable = [p for p in page_records if not p.get("skipped") and p.get("b64")]
        skipped_count = len(page_records) - len(usable)
        if not usable:
            raise CorruptFileError("PDF has no readable pages")
        if skipped_count:
            warnings.append(f"Skipped {skipped_count} blank page(s) in the PDF.")
        # Surface poor-quality warnings (deduped) before vision runs.
        seen_warnings: set[str] = set()
        for rec in usable:
            for w in (rec.get("quality") or {}).get("warnings", []):
                if w not in seen_warnings:
                    seen_warnings.add(w)
                    warnings.append(f"Page {rec['page_num']}: {w}")
        # Run all pages through vision in parallel — bounded concurrency.
        semaphore = asyncio.Semaphore(4)

        async def _read_page(rec):
            async with semaphore:
                try:
                    text = await _vision_extract_image(rec["b64"], label=f"pdf-page-{rec['page_num']}")
                    return rec["page_num"], text, None
                except Exception as e:
                    return rec["page_num"], "", str(e)

        results = await asyncio.gather(*[_read_page(r) for r in usable])
        results.sort(key=lambda x: x[0])
        chunks = []
        for page_num, page_text, err in results:
            if err:
                warnings.append(f"Vision read failed on page {page_num}: {err}")
            chunks.append(f"\n--- Page {page_num} ---\n{page_text}")
        return "\n".join(chunks), "pdf_scanned", page_count, warnings

    if ext in (".jpg", ".jpeg", ".png", ".heic", ".heif", ".webp"):
        try:
            b64, quality = _image_to_base64_jpeg(raw, ext)
        except CorruptFileError:
            raise
        # Surface quality warnings before vision runs.
        for w in quality.get("warnings", []):
            warnings.append(w)
        if quality.get("rating") == "poor":
            warnings.append(
                "Overall photo quality is poor. We'll do our best, but if results look wrong, "
                "try re-photographing in better light with the page flat and the camera focused."
            )
        text = await _vision_extract_image(b64, label=f"single-{ext.lstrip('.')}")
        if not text or len(text.strip()) < 50:
            warnings.append("Image read returned very little text. The photo may be too dark, blurry, or at an awkward angle.")
        return text, "image_vision", 1, warnings

    raise UnsupportedFormatError(ext)
